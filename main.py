import streamlit as st
import requests
import json
import re
import PyPDF2
from docx import Document
from fpdf import FPDF

# ==========================================
# 1. CONFIGURATION
# ==========================================
st.set_page_config(page_title="AI Resume Forge", layout="wide", page_icon="🎨")

if 'parsed_content' not in st.session_state:
    st.session_state['parsed_content'] = {}
if 'ats_analysis' not in st.session_state:
    st.session_state['ats_analysis'] = None

# ==========================================
# 2. CORE FUNCTIONS (UNCHANGED)
# ==========================================

def extract_text(file):
    try:
        if file.type == "application/pdf":
            reader = PyPDF2.PdfReader(file)
            text_list = []
            for page in reader.pages:
                t = page.extract_text()
                if t: text_list.append(t)
            full_text = " ".join(text_list)
            return full_text[:10000]
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            return full_text[:10000]
        return ""
    except Exception as e:
        return f"Error parsing file: {e}"

def clean_json_string(json_str):
    try:
        json_str = re.sub(r'```json\s*', '', json_str, flags=re.IGNORECASE)
        json_str = re.sub(r'```', '', json_str)
        start = json_str.find('{')
        end = json_str.rfind('}')
        if start != -1 and end != -1:
            return json_str[start:end+1]
        return json_str
    except:
        return "{}"

# --- SMART API CALL (AUTO-SWITCHING) ---
def call_gemini_api(prompt, api_key, primary_model):
    clean_key = api_key.strip()
    # Priority Chain: Selected -> 2.0 -> 1.5 (Safety Net)
    model_chain = [primary_model, "gemini-2.0-flash", "gemini-1.5-flash"]
    model_chain = list(dict.fromkeys(model_chain)) 
    
    last_error = ""

    for model in model_chain:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={clean_key}"
        headers = {'Content-Type': 'application/json'}
        data = {"contents": [{"parts": [{"text": prompt}]}]}
        
        try:
            response = requests.post(url, headers=headers, json=data, timeout=30)
            
            if response.status_code == 429: # Quota hit
                continue 
            
            if response.status_code != 200:
                return {"error": f"Error on {model}: {response.text}"}
                
            result = response.json()
            try:
                text = result['candidates'][0]['content']['parts'][0]['text']
                if model != primary_model:
                    text += f"\n\n(Note: Switched to {model} due to quota limits)"
                return text
            except:
                return {"error": "Unexpected API structure"}
                
        except Exception as e:
            last_error = str(e)
            continue

    return {"error": f"All models failed. Last error: {last_error}"}

# --- WRAPPERS ---
def score_resume(resume, jd, api_key, model):
    prompt = f"""
    You are an ATS Scanner. Evaluate RESUME vs JD.
    JD: {jd[:2000]}
    RESUME: {resume[:4000]}
    Output JSON: {{ "match_percentage": 50, "missing_keywords": ["a","b"], "reasoning": "x", "improvement_tips": ["y"] }}
    """
    raw_response = call_gemini_api(prompt, api_key, model)
    if isinstance(raw_response, dict) and "error" in raw_response: return raw_response
    return json.loads(clean_json_string(raw_response))

def enhance_resume(resume, jd, api_key, model):
    prompt = f"""
    Rewrite resume for JD. Output JSON.
    JD: {jd[:1000]}
    RESUME: {resume[:2000]}
    Output JSON: {{ "name": "A", "email": "B", "phone": "C", "summary": "D", "skills": ["E"], "experience": [{{ "title": "T", "company": "C", "points": ["P"] }}], "education": [{{ "degree": "D", "school": "S", "year": "Y" }}] }}
    """
    raw_response = call_gemini_api(prompt, api_key, model)
    if isinstance(raw_response, dict) and "error" in raw_response: return raw_response
    return json.loads(clean_json_string(raw_response))

# ==========================================
# 3. PDF GENERATOR (NEW)
# ==========================================
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        # No header content to keep it clean
        
    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.set_fill_color(200, 220, 255)
        self.cell(0, 6, title, 0, 1, 'L', 1)
        self.ln(4)

    def chapter_body(self, body):
        self.set_font('Arial', '', 11)
        self.multi_cell(0, 5, body)
        self.ln()

def generate_pdf(data):
    pdf = PDF()
    pdf.add_page()
    
    # Name & Contact
    pdf.set_font('Arial', 'B', 24)
    pdf.cell(0, 10, data.get('name', 'Name'), 0, 1, 'C')
    pdf.set_font('Arial', '', 12)
    pdf.cell(0, 10, f"{data.get('email', '')} | {data.get('phone', '')}", 0, 1, 'C')
    pdf.ln(10)
    
    # Summary
    pdf.chapter_title("Professional Summary")
    pdf.chapter_body(data.get('summary', ''))
    
    # Experience
    pdf.chapter_title("Work Experience")
    for job in data.get('experience', []):
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(0, 6, f"{job.get('title')} - {job.get('company')}", 0, 1)
        pdf.set_font('Arial', '', 11)
        for point in job.get('points', []):
             pdf.multi_cell(0, 5, f"- {point}")
        pdf.ln(3)

    # Education
    pdf.chapter_title("Education")
    for edu in data.get('education', []):
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(0, 6, f"{edu.get('degree')} - {edu.get('school')} ({edu.get('year')})", 0, 1)
    pdf.ln(5)
    
    # Skills
    pdf.chapter_title("Skills")
    pdf.chapter_body(", ".join(data.get('skills', [])))
    
    return pdf.output(dest='S').encode('latin1')

# ==========================================
# 4. LATEX TEMPLATES (NEW)
# ==========================================

TEMPLATES = {
    "Classic": r"""
\documentclass[a4paper,10pt]{article}
\usepackage[left=1in,right=1in,top=1in,bottom=1in]{geometry}
\usepackage{titlesec}
\titleformat{\section}{\large\bfseries\uppercase}{}{0em}{}[\titlerule]
\begin{document}
\begin{center}
    {\Huge \textbf{Var_Name}} \\
    Var_Email | Var_Phone
\end{center}
\section{Summary}
Var_Summary
\section{Experience}
Var_Experience
\section{Education}
Var_Education
\section{Skills}
Var_Skills
\end{document}
    """,
    
    "Modern": r"""
\documentclass[11pt,a4paper,sans]{moderncv}
\moderncvstyle{banking}
\moderncvcolor{blue}
\usepackage[scale=0.75]{geometry}
\name{Var}{Name}
\email{Var_Email}
\phone{Var_Phone}
\begin{document}
\makecvtitle
\section{Summary}
Var_Summary
\section{Experience}
Var_Experience
\section{Skills}
Var_Skills
\section{Education}
Var_Education
\end{document}
    """,
    
    "Technical": r"""
\documentclass[a4paper,10pt]{article}
\usepackage{geometry}
\geometry{a4paper, margin=0.75in}
\begin{document}
\noindent{\Huge \bfseries Var_Name} \\
\rule{\textwidth}{1pt}
Var_Email \hfill Var_Phone \\
\section*{TECHNICAL SKILLS}
Var_Skills
\section*{EXPERIENCE}
Var_Experience
\end{document}
    """
}

def fill_latex(template_name, data):
    # Simple string replacement to avoid Jinja dependency complexity for simple templates
    t = TEMPLATES[template_name]
    
    # Basic replacements
    t = t.replace("Var_Name", data.get('name', ''))
    t = t.replace("Var_Email", data.get('email', ''))
    t = t.replace("Var_Phone", data.get('phone', ''))
    t = t.replace("Var_Summary", data.get('summary', ''))
    
    # Format Lists
    exp_str = ""
    for job in data.get('experience', []):
        exp_str += f"\\textbf{{{job.get('title')}}} at \\textit{{{job.get('company')}}} \\\\ \n"
        exp_str += "\\begin{itemize} \n"
        for p in job.get('points', []):
            exp_str += f"\\item {p} \n"
        exp_str += "\\end{itemize} \\vspace{2mm} \n"
    t = t.replace("Var_Experience", exp_str)
    
    edu_str = ""
    for edu in data.get('education', []):
        edu_str += f"\\textbf{{{edu.get('degree')}}} - {edu.get('school')} ({edu.get('year')}) \\\\ \n"
    t = t.replace("Var_Education", edu_str)
    
    skills_str = ", ".join(data.get('skills', []))
    t = t.replace("Var_Skills", skills_str)
    
    return t

# ==========================================
# 5. USER INTERFACE
# ==========================================
st.title("🔄 AI Resume Builder")

with st.sidebar:
    st.header("Settings")
    api_key = st.text_input("Enter Gemini API Key", type="password")
    model = st.selectbox("Preferred Model:", ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-exp-1206"])
    st.info("ℹ️ System will auto-switch if quota limit is hit.")

st.header("1. Input")
input_method = st.radio("Input:", ["Upload", "Manual"], horizontal=True)
jd = st.text_area("Job Description")

raw_text = ""
if input_method == "Upload":
    f = st.file_uploader("File", type=['pdf', 'docx'])
    if f: raw_text = extract_text(f)
else:
    raw_text = st.text_area("Paste Text")

if raw_text and jd:
    st.header("2. Analysis")
    if st.button("Run ATS Check"):
        if not api_key: st.error("No Key")
        else:
            with st.spinner("Analyzing (Auto-Switching if needed)..."):
                analysis = score_resume(raw_text, jd, api_key, model)
                if "error" in analysis:
                    st.error(analysis["error"])
                else:
                    st.session_state['ats_analysis'] = analysis
                    st.session_state['parsed_content'] = analysis # Store temporarily
                    st.metric("Score", f"{analysis.get('match_percentage', 0)}%")
                    st.write(analysis.get("reasoning"))
                    st.write("Missing: " + ", ".join(analysis.get("missing_keywords", [])))
    
    st.header("3. Enhance")
    if st.button("Enhance"):
        if not api_key: st.error("No Key")
        else:
            with st.spinner("Enhancing..."):
                data = enhance_resume(raw_text, jd, api_key, model)
                if "error" in data: st.error(data["error"])
                else: st.session_state['data'] = data; st.json(data)

# ==========================================
# STEP 4: ENHANCED EXPORT (NEW)
# ==========================================
if st.session_state.get('data'):
    st.divider()
    st.header("4. Export Options")
    data = st.session_state['data']
    
    col_pdf, col_latex, col_word = st.columns(3)
    
    # 1. DIRECT PDF DOWNLOAD
    with col_pdf:
        st.subheader("📄 PDF")
        try:
            pdf_bytes = generate_pdf(data)
            st.download_button(
                label="Download PDF",
                data=pdf_bytes,
                file_name="resume.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.error(f"PDF Error: {e}")
            st.caption("Ensure 'fpdf' is installed.")

    # 2. LATEX DOWNLOAD
    with col_latex:
        st.subheader("📜 LaTeX")
        template_choice = st.selectbox("Select Template", ["Classic", "Modern", "Technical"])
        latex_code = fill_latex(template_choice, data)
        st.download_button(
            label="Download .tex",
            data=latex_code,
            file_name="resume.tex",
            mime="text/plain"
        )
        st.caption("Upload this .tex file to Overleaf.com for professional rendering.")

    # 3. WORD DOWNLOAD
    with col_word:
        st.subheader("📝 Word")
        doc = Document()
        doc.add_heading(data.get('name', 'Name'), 0)
        doc.add_paragraph(f"{data.get('email','')} | {data.get('phone','')}")
        doc.add_heading("Summary", 1)
        doc.add_paragraph(data.get('summary', ''))
        doc.add_heading("Experience", 1)
        for job in data.get('experience', []):
            doc.add_paragraph(f"{job.get('title')} - {job.get('company')}")
            for p in job.get('points', []):
                doc.add_paragraph(p, style='List Bullet')
        doc.save("resume.docx")
        with open("resume.docx", "rb") as f:
            st.download_button("Download .docx", f, "resume.docx")